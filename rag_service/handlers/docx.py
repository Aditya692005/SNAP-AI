from pathlib import Path

from .base import Parsed, Table


def parse(path: Path) -> Parsed:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]

    tables = []
    for idx, table in enumerate(doc.tables):
        header = None
        rows = []
        for r_i, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if r_i == 0:
                header = cells
            elif any(cells):
                if header and len(header) == len(cells):
                    rows.append(dict(zip(header, cells)))
                else:
                    rows.append({f"col_{i + 1}": v for i, v in enumerate(cells)})
        if rows:
            tables.append(
                Table(
                    rows=rows,
                    table_name=f"Table {idx + 1}",
                    heading_context=", ".join(header) if header else None,
                    table_index=idx,
                )
            )

    text_chunks = ["\n".join(paragraphs)] if paragraphs else []
    return Parsed(text_chunks=text_chunks, tables=tables)
